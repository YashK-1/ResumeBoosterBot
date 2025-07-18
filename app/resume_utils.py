from docx import Document
from docx.shared import Pt

def write_resume_to_docx(resume_text: str, output_path: str):
    """
    Writes the improved resume text to a .docx file.

    Parameters:
    - resume_text (str): The AI-generated or refined resume content.
    - output_path (str): Full path to save the .docx file.
    """
    try:
        # Create a new Word document
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Split into lines and write to document
        for line in resume_text.strip().split('\n'):
            if line.strip() == "":
                doc.add_paragraph("")  # Add a blank line
            elif line.strip().endswith(":"):
                # Section heading
                doc.add_paragraph(line.strip(), style='Heading 2')
            else:
                doc.add_paragraph(line.strip())

        # Save the document
        doc.save(output_path)

    except Exception as e:
        raise RuntimeError(f"❌ Failed to write .docx: {e}")
